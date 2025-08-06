import os
import io
import json
import re
import datetime
from typing import List, Dict, Tuple

import streamlit as st
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from openai import OpenAI

# ---------- Load env & page ----------
load_dotenv()
st.set_page_config(page_title="AI Resume Editor (MVP)", page_icon="🧠", layout="wide")
st.title("🧠 AI Resume Editor — MVP")

# ---------- Prompts & Firebase ----------
from prompts import (
    SYSTEM_PROMPT,
    ANALYZE_PROMPT,
    REWRITE_BULLETS_PROMPT,
    REWRITE_SUMMARY_PROMPT,
    COVER_LETTER_PROMPT,       # existing (cover letter)
    GPT_MATCH_SCORE_PROMPT,    # NEW (Phase 2.5 – semantic scoring)
)
from firebase_utils import init_firebase, save_history, get_recent_history

# ---------- Sidebar controls ----------
st.sidebar.header("Setup")
api_key_input = st.sidebar.text_input("OpenAI API Key (optional if using .env)", type="password")
model = st.sidebar.text_input("Model", value=os.getenv("MODEL", "gpt-4o-mini"))
tone = st.sidebar.selectbox("Tone", ["neutral", "impactful", "conservative", "friendly"], index=1)
seniority = st.sidebar.selectbox("Seniority", ["entry-level", "mid-level", "senior"], index=1)

# ---------- OpenAI client (cached) ----------
@st.cache_resource(show_spinner=False)
def _get_client(key: str):
    if not key:
        raise RuntimeError("OpenAI API key missing.")
    return OpenAI(api_key=key)

def get_client() -> OpenAI:
    key = api_key_input or os.getenv("OPENAI_API_KEY")
    if not key:
        st.error("Add your OpenAI API Key in the sidebar or .env file.")
        st.stop()
    return _get_client(key)

# ---------- File parsers ----------
def read_text_from_pdf(file) -> str:
    try:
        reader = PdfReader(file)
        pages = [p.extract_text() or "" for p in reader.pages]
        return "\n".join(pages).strip()
    except Exception as e:
        st.warning(f"PDF parsing error: {e}")
        return ""

def read_text_from_docx(file) -> str:
    try:
        doc = Document(file)
        return "\n".join([(p.text or "") for p in doc.paragraphs]).strip()
    except Exception as e:
        st.warning(f"DOCX parsing error: {e}")
        return ""

# ---------- Theme (CSS injector) ----------
def inject_theme_css():
    """Force the dark theme styles (no light mode)."""
    c = {
        "primary": "#60A5FA",
        "bg": "#0b1220",
        "text": "#e5e7eb",
        "surface": "#121a2b",
        "muted": "#94a3b8",
        "border": "#1f2a44",
    }
    st.markdown(
        f"""
        <style>
          :root {{
            --primary-color: {c['primary']};
            --app-bg: {c['bg']};
            --text-color: {c['text']};
            --surface: {c['surface']};
            --muted: {c['muted']};
            --border: {c['border']};
          }}
          .stApp {{
            background: var(--app-bg);
            color: var(--text-color);
          }}
          /* Sidebar */
          [data-testid="stSidebar"] {{
            background: linear-gradient(180deg, rgba(0,0,0,0.18), rgba(0,0,0,0)) , var(--surface);
            border-right: 1px solid var(--border);
          }}
          /* Buttons */
          .stButton>button, .stDownloadButton>button {{
            border-radius: 10px;
            font-weight: 600;
            border: 1px solid var(--border);
            background: var(--surface);
            color: var(--text-color);
          }}
          .stButton>button:hover, .stDownloadButton>button:hover {{
            transform: translateY(-1px);
            box-shadow: 0 6px 14px rgba(0,0,0,0.35);
          }}
          /* Inputs */
          .stTextArea textarea, .stTextInput input, div[data-baseweb="select"] > div {{
            background: var(--surface) !important;
            color: var(--text-color) !important;
            border: 1px solid var(--border) !important;
            border-radius: 10px !important;
          }}
          /* Progress bar primary color */
          .stProgress > div > div > div > div {{
            background-color: var(--primary-color);
          }}
          /* Headers & captions */
          h1, h2, h3, h4, h5, h6, .stCaption, .caption {{
            color: var(--text-color);
          }}
          /* Chips */
          .score-chip {{
            display:inline-block; padding:4px 10px; border-radius:999px; font-weight:600;
            border:1px solid var(--border); background: var(--surface);
          }}
        </style>
        """,
        unsafe_allow_html=True,
    )

# >>> Force dark mode globally
inject_theme_css()

# ---------- Helpers ----------
def make_docx(summary: str, bullets: List[str], extra_sections: Dict[str, List[str]] | None = None) -> bytes:
    doc = Document()
    doc.add_heading("Resume", level=0)

    if summary:
        doc.add_heading("Summary", level=1)
        for line in summary.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    if bullets:
        doc.add_heading("Experience Highlights", level=1)
        for b in bullets:
            if b.strip():
                doc.add_paragraph(b.strip(), style="List Bullet")

    if extra_sections:
        for name, items in extra_sections.items():
            doc.add_heading(name, level=1)
            for it in items:
                if it.strip():
                    doc.add_paragraph(it.strip(), style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def safe_load_json(text: str) -> dict:
    try:
        return json.loads(text)
    except Exception:
        pass
    m = re.search(r"\{.*\}", text, re.S)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return {"notes": [text]}
    return {"notes": [text]}

def pct(n: int, d: int) -> int:
    return int(round(100 * n / d)) if d else 0

# ---------- Keyword & TF‑IDF helpers (pure‑Python, no sklearn) ----------
import math
import string
from collections import Counter

_STOPWORDS = {
    "a","an","the","and","or","but","if","then","so","of","for","to","from","in",
    "on","at","by","with","as","is","are","was","were","be","been","being",
    "this","that","these","those","it","its","i","you","he","she","we","they",
    "my","your","our","their","me","him","her","us","them","will","can","could",
    "should","would","may","might","do","does","did","done","over","into","per"
}

def _normalize(text: str) -> str:
    text = text.lower()
    allowed = set("+#")  # keep for C++ / C#
    table = str.maketrans({ch: " " for ch in string.punctuation if ch not in allowed})
    return text.translate(table)

def _tokens(text: str) -> list[str]:
    text = _normalize(text)
    toks = [t for t in (w.strip() for w in text.split())
            if t and t not in _STOPWORDS and not t.isdigit() and len(t) >= 2]
    # add bigrams (like sklearn ngram_range=(1,2))
    bigrams = [f"{toks[i]} {toks[i+1]}" for i in range(len(toks)-1)]
    return toks + bigrams

def extract_keywords(text: str) -> list[str]:
    # unigrams only for coverage display
    text = _normalize(text)
    toks = [t for t in (w.strip() for w in text.split())
            if t and t not in _STOPWORDS and not t.isdigit() and len(t) >= 2]
    return toks

def coverage(cv_keys: list[str], jd_keys: list[str]) -> tuple[list[str], list[str]]:
    cv_set = set(cv_keys)
    jd_set = set(jd_keys)
    covered = sorted(jd_set & cv_set)
    missing = sorted(jd_set - cv_set)
    return covered, missing

def _tfidf_vector(doc_tokens: list[str], df: dict[str, int], N_docs: int) -> dict[str, float]:
    tf = Counter(doc_tokens)
    vec = {}
    L = max(1, len(doc_tokens))
    for term, f in tf.items():
        dfi = df.get(term, 0)
        # smoothed idf: log((1+N)/(1+df)) + 1
        idf = math.log((1 + N_docs) / (1 + dfi)) + 1.0
        vec[term] = (f / L) * idf
    return vec

def _cosine_sparse(v1: dict[str, float], v2: dict[str, float]) -> float:
    dot = 0.0
    for k, a in v1.items():
        b = v2.get(k)
        if b is not None:
            dot += a * b
    n1 = math.sqrt(sum(a*a for a in v1.values()))
    n2 = math.sqrt(sum(b*b for b in v2.values()))
    if n1 == 0.0 or n2 == 0.0:
        return 0.0
    return dot / (n1 * n2)

def tfidf_match_score(resume_text: str, job_text: str) -> tuple[int, dict]:
    """
    Pure‑Python TF‑IDF cosine similarity on unigrams+bigrams.
    Returns (score_0_100, {"top_jd_terms": [...]})
    """
    if not resume_text.strip() or not job_text.strip():
        return 0, {"top_jd_terms": []}

    jd_toks = _tokens(job_text)
    cv_toks = _tokens(resume_text)

    # document frequencies over 2 docs
    df = Counter(set(jd_toks))
    for t in set(cv_toks):
        df[t] += 1
    N = 2

    jd_vec = _tfidf_vector(jd_toks, df, N)
    cv_vec = _tfidf_vector(cv_toks, df, N)

    sim = max(0.0, min(1.0, _cosine_sparse(jd_vec, cv_vec)))
    score = int(round(sim * 100))

    top_terms = [w for w, _ in sorted(jd_vec.items(), key=lambda x: x[1], reverse=True)[:15]]
    return score, {"top_jd_terms": top_terms}

# ---------- (Removed) Appearance toggle ----------
# No sidebar toggle; dark mode is always on via inject_theme_css()

# ---------- Shared Input (visible to both tabs) ----------
with st.expander("1) Upload Resume & Paste Job Description", expanded=True):
    resume_file = st.file_uploader("Resume (.pdf, .docx, or .txt)", type=["pdf", "docx", "txt"])
    job_text = st.text_area("Job Description", height=220, placeholder="Paste the target job description here")

    resume_text = ""
    if resume_file:
        name = (resume_file.name or "").lower()
        if name.endswith(".pdf"):
            resume_text = read_text_from_pdf(resume_file)
        elif name.endswith(".docx"):
            resume_text = read_text_from_docx(resume_file)
        else:
            resume_text = resume_file.read().decode("utf-8", errors="ignore").strip()

    col_a, col_b = st.columns(2)
    with col_a:
        st.markdown("**Resume (parsed)**")
        st.text_area(" ", value=resume_text, height=250, key="resume_preview")
    with col_b:
        st.markdown("**Job Description**")
        st.text_area("  ", value=job_text, height=250, key="jd_preview")

# Tabs hold the two workflows side-by-side in the same UI
tab1, tab2 = st.tabs(["📄 Match & Tailor Resume", "💌 Cover Letter Generator"])

# ---------- State placeholders ----------
analysis_json: dict = {}
rewritten_bullets: List[str] = []
summary_text: str = ""
session_saved = False

# =======================================
# TAB 1: Resume Analysis / Rewrite / Generate
# =======================================
with tab1:
    # ---------- Actions ----------
    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
    with col_btn1:
        analyze_clicked = st.button("🔎 Analyze Fit")
    with col_btn2:
        rewrite_clicked = st.button("✍️ Rewrite Bullets")
    with col_btn3:
        generate_clicked = st.button("📄 Generate Tailored Resume")

    # ---------- Analyze Fit (Job Match) ----------
    if analyze_clicked:
        if not resume_text or not job_text:
            st.warning("Please add both resume and job description.")
        else:
            client = get_client()
            with st.spinner("Analyzing…"):
                prompt = ANALYZE_PROMPT.format(
                    resume_text=resume_text[:15000],
                    job_text=job_text[:6000]
                )
                resp = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.2,
                )
                raw = resp.choices[0].message.content
                analysis_json = safe_load_json(raw)

            # --- NEW: Semantic (GPT) match score ---
            gpt_prompt = GPT_MATCH_SCORE_PROMPT.format(
                resume_text=resume_text[:15000],
                job_text=job_text[:6000],
            )
            resp2 = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": gpt_prompt},
                ],
                temperature=0.0,
            )
            gpt_raw = resp2.choices[0].message.content or "{}"
            gpt_json = safe_load_json(gpt_raw)

            st.subheader("🔍 Job Match")

            # A) TF‑IDF score (primary baseline)
            tfidf_score, tfidf_details = tfidf_match_score(resume_text, job_text)

            # B) Keyword coverage (transparent baseline)
            jd_keys = extract_keywords(job_text)
            cv_keys = extract_keywords(resume_text)
            covered, missing = coverage(cv_keys, jd_keys)
            naive_score = pct(len(covered), len(covered) + len(missing))

            # C) Model-provided score (if present)
            model_score = analysis_json.get("overall_fit") or analysis_json.get("match_score")
            try:
                model_score = int(float(model_score))
            except Exception:
                model_score = None

            # D) NEW: GPT semantic score
            gpt_score = gpt_json.get("match_score")
            try:
                gpt_score = int(float(gpt_score))
            except Exception:
                gpt_score = None

            # E) NEW: Hybrid score (semantic-forward, transparent baselines)
            weights = {
                "gpt": 0.6 if isinstance(gpt_score, int) else 0.0,
                "tfidf": 0.3,
                "naive": 0.1,
            }
            hybrid_score = round(
                (weights["gpt"] * (gpt_score or 0)) +
                (weights["tfidf"] * (tfidf_score or 0)) +
                (weights["naive"] * (naive_score or 0))
            )

            # Display preference: Hybrid > model > TF‑IDF > naive
            display_score = (
                hybrid_score if hybrid_score > 0
                else (model_score if isinstance(model_score, int)
                else (tfidf_score if tfidf_score else naive_score))
            )

            # --- Colored score chip ---
            score_color = (
                "#ef4444" if display_score < 40 else  # red
                "#f59e0b" if display_score < 70 else  # amber
                "#22c55e"                            # green
            )
            st.markdown(
                f"**Match Score:** <span class='score-chip' style='background:{score_color}; color:white;'>{display_score}%</span>",
                unsafe_allow_html=True
            )
            st.progress(min(max(display_score, 0), 100) / 100)

            cols = st.columns(4)
            with cols[0]:
                st.caption("Hybrid score (GPT + TF‑IDF)")
                st.write(f"{hybrid_score}")
            with cols[1]:
                st.caption("GPT semantic score")
                st.write(f"{(str(gpt_score) + '%') if gpt_score is not None else '—'}")
            with cols[2]:
                st.caption("TF‑IDF score")
                st.write(f"{tfidf_score}%")
            with cols[3]:
                st.caption("Keyword coverage (naive)")
                st.write(f"{naive_score}%")

            # Reasons & Missing keywords
            cols2 = st.columns(2)
            with cols2[0]:
                st.markdown("**Top reasons (model):**")
                for r in analysis_json.get("notes", [])[:5]:
                    st.write("• " + str(r))
            with cols2[1]:
                missing_kw = analysis_json.get("missing_keywords") or list(sorted(set(missing)))
                if missing_kw:
                    st.markdown("**Missing keywords:**")
                    st.write(", ".join(sorted({str(x).lower().strip() for x in missing_kw})[:30]))
                else:
                    st.markdown("**Missing keywords:** None 🎉")

            # NEW: GPT rationale & matched skills
            st.markdown("**GPT semantic assessment (why this score):**")
            if isinstance(gpt_json, dict) and gpt_json.get("summary_reason"):
                st.write(gpt_json["summary_reason"])
            if isinstance(gpt_json, dict) and gpt_json.get("skills_matched"):
                st.caption("Skills matched (GPT)")
                st.write(", ".join([str(s) for s in gpt_json.get("skills_matched", [])][:15]))

            # Suggested skills & TF‑IDF top terms
            if analysis_json.get("skills_to_surface"):
                st.markdown("**Skills to surface (model):**")
                st.write(", ".join(analysis_json["skills_to_surface"][:12]))

            st.markdown("**Top JD terms by TF‑IDF (signals to mirror naturally):**")
            st.write(", ".join(tfidf_details.get("top_jd_terms", [])[:15]))

            st.markdown("**Keyword Coverage (naive baseline):**")
            st.write({"covered": covered[:60], "missing": missing[:60]})

            # Save history
            try:
                save_history({
                    "action": "analyze_fit",
                    "resume_text": resume_text[:12000],
                    "job_description": job_text[:8000],
                    "analysis_json": analysis_json,
                    "scores": {
                        "model": model_score,
                        "tfidf": tfidf_score,
                        "naive": naive_score,
                        "gpt": gpt_score,
                        "hybrid": hybrid_score,
                    },
                    "gpt_assessment": {
                        "summary_reason": (gpt_json.get("summary_reason") if isinstance(gpt_json, dict) else None),
                        "skills_matched": (gpt_json.get("skills_matched", []) if isinstance(gpt_json, dict) else []),
                    },
                    "tfidf_top_jd_terms": tfidf_details.get("top_jd_terms", []),
                })
                session_saved = True
            except Exception as e:
                st.info(f"(Skipping save: {e})")

    # ---------- Rewrite Bullets ----------
    if rewrite_clicked:
        if not resume_text or not job_text:
            st.warning("Please add both resume and job description.")
        else:
            client = get_client()
            with st.spinner("Rewriting bullets…"):
                prompt = REWRITE_BULLETS_PROMPT.format(
                    resume_text=resume_text[:15000],
                    job_text=job_text[:6000],
                    tone=tone,
                    seniority=seniority,
                )
                resp = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.4,
                )
                raw = resp.choices[0].message.content
                try:
                    rewritten_bullets = json.loads(raw).get("bullets", [])
                except Exception:
                    rewritten_bullets = [b.strip("-• \n") for b in re.findall(r"[-•]\s*(.+)", raw)]

            st.subheader("✍️ Suggested Bullets")
            if not rewritten_bullets:
                st.info("No bullets returned. Try adjusting tone/seniority or re-run.")
            for b in rewritten_bullets:
                st.write("• " + b)

            # Save history
            try:
                save_history({
                    "action": "rewrite_bullets",
                    "resume_text": resume_text[:12000],
                    "job_description": job_text[:8000],
                    "bullets": rewritten_bullets[:20],
                    "tone": tone,
                    "seniority": seniority,
                })
                session_saved = True
            except Exception as e:
                st.info(f"(Skipping save: {e})")

    # ---------- Generate Tailored Resume ----------
    if generate_clicked:
        if not resume_text or not job_text:
            st.warning("Please add both resume and job description.")
        else:
            client = get_client()

            # 1) Tailored summary
            with st.spinner("Generating tailored summary…"):
                prompt = REWRITE_SUMMARY_PROMPT.format(
                    resume_text=resume_text[:15000],
                    job_text=job_text[:6000],
                    tone=tone,
                    seniority=seniority,
                )
                resp = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.4,
                )
                summary_text = resp.choices[0].message.content.strip()

            # 2) Bullets (if user didn’t create them already)
            if not rewritten_bullets:
                with st.spinner("Generating bullets…"):
                    prompt = REWRITE_BULLETS_PROMPT.format(
                        resume_text=resume_text[:15000],
                        job_text=job_text[:6000],
                        tone=tone,
                        seniority=seniority,
                    )
                    resp = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": SYSTEM_PROMPT},
                            {"role": "user", "content": prompt},
                        ],
                        temperature=0.4,
                    )
                    raw = resp.choices[0].message.content
                    try:
                        rb = json.loads(raw).get("bullets", [])
                    except Exception:
                        rb = [b.strip("-• \n") for b in re.findall(r"[-•]\s*(.+)", raw)]
                    rewritten_bullets.extend(rb)

            # 3) Build DOCX
            key_skills = list(sorted(set(extract_keywords(job_text))))[:12]
            doc_bytes = make_docx(summary_text, rewritten_bullets[:8], {"Key Skills": key_skills})

            st.download_button(
                "Download Tailored Resume (DOCX)",
                data=doc_bytes,
                file_name="tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            # Save history
            try:
                save_history({
                    "action": "generate_resume",
                    "resume_text": resume_text[:12000],
                    "job_description": job_text[:8000],
                    "summary": summary_text,
                    "bullets": rewritten_bullets[:20],
                    "tone": tone,
                    "seniority": seniority,
                    "key_skills": key_skills,
                })
                session_saved = True
            except Exception as e:
                st.info(f"(Skipping save: {e})")

# =======================================
# TAB 2: Cover Letter Generator
# =======================================
with tab2:
    st.header("💌 Cover Letter Generator")
    st.caption("Uses your uploaded resume and pasted job description above. You can tweak text below before generating.")

    # Prefill with parsed text so user can tweak
    resume_cl = st.text_area("Resume text for cover letter", value=resume_text, height=250, key="cl_resume")
    job_cl = st.text_area("Job description for cover letter", value=job_text, height=250, key="cl_jd")

    if st.button("✉️ Generate Cover Letter"):
        if not resume_cl or not job_cl:
            st.warning("Please provide both resume text and job description.")
        else:
            client = get_client()
            prompt = COVER_LETTER_PROMPT.format(
                resume_text=resume_cl[:15000],
                job_text=job_cl[:6000],
                tone=tone,
                seniority=seniority,
            )
            with st.spinner("Writing cover letter…"):
                resp = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": SYSTEM_PROMPT},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=0.5,
                )
                cover_letter = resp.choices[0].message.content.strip()

            st.subheader("✉️ Your Cover Letter")
            st.write(cover_letter)

            st.download_button(
                "Download Cover Letter (.txt)",
                data=cover_letter.encode("utf-8"),
                file_name="cover_letter.txt",
                mime="text/plain",
            )

            # Save to Firebase
            try:
                save_history({
                    "action": "generate_cover_letter",
                    "resume_text": resume_cl[:12000],
                    "job_description": job_cl[:8000],
                    "cover_letter": cover_letter,
                    "tone": tone,
                    "seniority": seniority,
                })
                st.success("Saved to history.")
            except Exception as e:
                st.info(f"(Skipping save: {e})")

# ---------- History viewer ----------
st.markdown("---")
with st.expander("🕒 Recent History (last 10)"):
    try:
        records = get_recent_history(limit=10)
        if not records:
            st.write("No history yet.")
        else:
            for r in records:
                ts = r.get("timestamp", "—")
                act = r.get("action", "—")
                st.markdown(f"**{ts}** — _{act}_")
                if "scores" in r:
                    st.write(r["scores"])
                if act == "generate_resume":
                    st.caption("Summary")
                    st.write(r.get("summary", "")[:700])
                    st.caption("Bullets")
                    for b in (r.get("bullets") or [])[:6]:
                        st.write("• " + str(b))
                if act == "analyze_fit" and r.get("analysis_json"):
                    aj = r["analysis_json"]
                    if aj.get("missing_keywords"):
                        st.caption("Missing keywords (model)")
                        st.write(", ".join(aj["missing_keywords"][:20]))
                if act == "generate_cover_letter":
                    st.caption("Cover Letter (preview)")
                    st.write((r.get("cover_letter") or "")[:700])
    except Exception as e:
        st.info(f"(History unavailable: {e})")

# ---------- Footer ----------

st.markdown(
    """
---
👨‍💻 Built with: [Streamlit](https://streamlit.io), [OpenAI API](https://platform.openai.com), [Firebase](https://firebase.google.com)  
📁 Project: [resume-ai-editor](https://github.com/Harsha-03/resume-ai-editor)  
💬 Feedback & collaboration welcome! → harshaasapu.b@gmail.com
"""
)

if 'session_saved' in locals() and session_saved:
    st.success("Session saved to Firestore.")

